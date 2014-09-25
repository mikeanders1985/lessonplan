from docx import *
from docx.shared import Inches
from docx.shared import Cm


class LessonPlanCreator(object):
    
    def __init__(self, name, time, filename, template):
        
        # Constructor function. Creates variables for text entry.
        self.name = name
        self.time = time
        self.template = template
        self.document = Document(self.template)
        self.filename = filename
        
        # Contains dictionaries for first page which allow text entry.
        self.first_row = {0 : ('Target from previous lesson:\n%s\n%s' % ('', '')), 1: 'Evidence towards Standards:'}
        self.second_row = {0: ('Date:\n%s' % ''), 1: 'Time:\n%s' % self.time, 2: 'Class:\n%s' % self.name, 3: 'Attainment Range:', 4: 'Number of pupils', 5: 'M:', 6: 'F:'}
        self.third_row = {0: 'Topic:\n%s' % ''}
        self.fourth_row = {0: ('Prior Learning (T2/T3):\n%s \n%s' % ('','')), 1: 'Potential Misconceptions (T3):'}
        self.fifth_row = {0: 'Learning Objectives  / Learning Focus (T3,4)\n(What are pupils learning to do?)', 1: 'Learning Outcomes / Success Criteria (T1, 2, 6)\n(What will pupils be able to do or demonstrate?)'}
        self.sixth_row = {0: '\n%s\n%s\n%s\n%s\n%s\n%s' % ('', '', '', '', '', ''), 1: ''}
        self.seventh_row = {0: 'Literacy and Oracy(T3):\n%s\n%s' % ('', ''), 1: 'Numeracy (T3):', 2: 'Technology/Cross curricular themes:'}
        self.eighth_row = {0: 'Meeting the Needs of All Learners (T5)  (consider SEN, G&T, EAL, PP/FSM Behaviour, low literacy/numeracy etc.)'}
        self.ninth_row = {0: 'Consideration', 1: 'Pupils', 2: 'Strategy'}
        self.tenth_row = {0: '\n%s\n%s' % ('', ''), 1: '', 2: ''}
        self.eleventh_row = {0: 'TASK', 1: 'Completely Different Task or Worksheet\n%s' % '', 2: 'Open-ended Task\n%s' % '', 3: 'Individual Work\n%s' % '', 4: 'Self-Supported Study\n%s' % '', 5: 'VAK\n%s' % ''}
        self.twelfth_row = {0: 'RESOURCES', 1: 'Work in Special Format\n%s' % '', 2: 'ICT/technology\n%s' % '', 3: 'Additional Resources\n%s' % '', 4: 'Homework challenge for G&T\n%s' % '', 5: 'Extension\n%s' % ''}
        self.thirteenth_row = {0: 'LITERACY', 1: 'Work Format\n%s' % '', 2: 'Cloze Procedure\n%s' % '', 3: 'Writing Frames / scaffolding\n%s' % '', 4: 'Reading/writing support\n%s' % '', 5: 'Numeracy support\n%s' % ''}
        self.fourteenth_row = {0: 'SUPPORT', 1: 'Use of Classroom Assistant (T8)\n%s' % '', 2: 'Group Work\n%s' % '', 3: 'Peer Support/ Collaboration\n%s' % '', 4: 'Intervention\n%s' % '', 5: 'Increasing Levels of Difficulty\n%s' % ''}
        self.fifteenth_row = {0: 'QUESTIONING & OUTCOME', 1: 'Directed questioning', 2: 'Use of Target Questions', 3: 'Outcome', 4: 'Accountability/engagement', 5: 'Other'}
        
        # Contains dictionaries which allow text entry for the second page.
        self.page_two_first_row = {0: 'Homework Activity (T4):\n%s\n%s' % ('', ''), 1: 'Completion Date:'}
        self.page_two_second_row = {0: '', 1: 'Time', 2: 'Teacher Activity (T1, 3, 4, 7)\nUSE OF SUPPORT ASSISTANT (T8)', 3: 'Pupil Learning Activities (T4, T5)', 4: 'Pupil Progress (T2) and Assessment (T6)', 5: 'Evaluation & Reflection'}
        self.page_two_third_row = {0: '\n%s\n%s\n%s\n%s\n%s\n%s\n%s' % ('S', 't', 'a', 'r', 't', 'e', 'r'), 1: '', 2: '', 3: '', 4: '', 5: ''}
        self.page_two_fourth_row = {0: '\n%s\n%s\n%s\n%s\n%s\n%s' % ('', 'M', 'a', 'i', 'n', ''), 1: '', 2: '', 3: '', 4: '', 5: ''}
        self.page_two_fifth_row = {0: '\n%s\n%s\n%s\n%s\n%s\n%s\n%s' % ('P', 'l', 'e', 'n', 'a', 'r', 'y'), 1: '', 2: '', 3: '', 4: '', 5: ''}
    
    def page_one_creator(self):
        
        # Creates tables for the first page.    
        self.first_row_table = self.document.add_table(rows = 1, cols = 2, style = 'firstrow')
        self.second_row_table = self.document.add_table(rows = 1, cols = 7, style = 'secondrow')
        self.third_row_table = self.document.add_table(rows = 1, cols = 1, style = 'thirdrow')
        self.fourth_row_table = self.document.add_table(rows = 1, cols = 2, style = 'fourthrow')
        self.fifth_row_table = self.document.add_table(rows = 1, cols = 2, style = 'fifthrow')
        self.sixth_row_table = self.document.add_table(rows = 1, cols = 2, style = 'sixthrow')
        self.seventh_row_table = self.document.add_table(rows = 1, cols = 3, style = 'thirdrow')
        self.eighth_row_table = self.document.add_table(rows = 1, cols = 1, style = 'eighthrow')
        self.ninth_row_table = self.document.add_table(rows = 1, cols = 3, style = 'ninthrow')
        self.tenth_row_table = self.document.add_table(rows = 1, cols = 3, style = 'sixthrow')
        self.eleventh_row_table = self.document.add_table(rows = 1, cols = 6, style = 'eleventhrow')
        self.twelfth_row_table = self.document.add_table(rows = 1, cols = 6, style = 'eleventhrow')
        self.thirteenth_row_table = self.document.add_table(rows = 1, cols = 6, style = 'eleventhrow')
        self.fourteenth_row_table = self.document.add_table(rows = 1, cols = 6, style = 'eleventhrow')
        self.fifteenth_row_table = self.document.add_table(rows = 1, cols = 6, style = 'fifteenthrow')
        
        # Page break to allow second page to be added
        self.document.add_page_break()
        
    def page_two_creator(self):
        
        # Creates tables for the second page.
        self.page_two_first_row_table = self.document.add_table(rows = 1, cols = 2, style = 'thirdrow')
        self.page_two_second_row_table = self.document.add_table(rows = 1, cols = 6, style = 'p2secondrow')
        self.page_two_third_row_table = self.document.add_table(rows = 1, cols = 6, style = 'p2thirdrow')
        self.page_two_fourth_row_table = self.document.add_table(rows = 1, cols = 6, style = 'p2fourthrow')
        self.page_two_fifth_row_table = self.document.add_table(rows = 1, cols = 6, style = 'p2fifthrow')

    def page_one_fill(self):
        
        # Fills tables for the first page, depending on input.
        self.first_row_cells = self.first_row_table.rows[0].cells
        self.first_row_cells[0].text = self.first_row[0]
        self.first_row_cells[0].width = Cm(12.74)
        self.first_row_cells[1].text = self.first_row[1]
        self.first_row_cells[1].width = Cm(11.63)
        
        self.second_row_cells = self.second_row_table.rows[0].cells
        self.second_row_cells[0].text = self.second_row[0]
        self.second_row_cells[0].width = Cm(4.11)
        self.second_row_cells[1].text = self.second_row[1]
        self.second_row_cells[1].width = Cm(4.44)
        self.second_row_cells[2].text = self.second_row[2]
        self.second_row_cells[2].width = Cm(4.39)
        self.second_row_cells[3].text = self.second_row[3]
        self.second_row_cells[3].width = Cm(5.03)
        self.second_row_cells[4].text = self.second_row[4]
        self.second_row_cells[4].width = Cm(3.75)
        self.second_row_cells[5].text = self.second_row[5]
        self.second_row_cells[5].width = Cm(3.0)
        self.second_row_cells[6].text = self.second_row[6]
        self.second_row_cells[6].width = Cm(3.21)
        
        self.third_row_cells = self.third_row_table.rows[0].cells
        self.third_row_cells[0].text = self.third_row[0]
        self.third_row_cells[0].width = Cm(27.91)
        
        self.fourth_row_cells = self.fourth_row_table.rows[0].cells
        self.fourth_row_cells[0].text = self.fourth_row[0]
        self.fourth_row_cells[0].width = Cm(12.91)
        self.fourth_row_cells[1].text = self.fourth_row[1]
        self.fourth_row_cells[1].width = Cm(14.99)
        
        self.fifth_row_cells = self.fifth_row_table.rows[0].cells
        self.fifth_row_cells[0].text = self.fifth_row[0]
        self.fifth_row_cells[0].width = Cm(12.74)
        self.fifth_row_cells[1].text = self.fifth_row[1]
        self.fifth_row_cells[1].width = Cm(14.8)
        
        self.sixth_row_cells = self.sixth_row_table.rows[0].cells
        self.sixth_row_cells[0].text = self.sixth_row[0]
        self.sixth_row_cells[0].width = Cm(12.74)
        self.sixth_row_cells[1].text = self.sixth_row[1]
        self.sixth_row_cells[1].width = Cm(14.8)
        
        self.seventh_row_cells = self.seventh_row_table.rows[0].cells
        self.seventh_row_cells[0].text = self.seventh_row[0]
        self.seventh_row_cells[0].width = Cm(9.18)
        self.seventh_row_cells[1].text = self.seventh_row[1]
        self.seventh_row_cells[1].width = Cm(9.18)
        self.seventh_row_cells[2].text = self.seventh_row[2]
        self.seventh_row_cells[2].width = Cm(9.19)
        
        self.eighth_row_cells = self.eighth_row_table.rows[0].cells
        self.eighth_row_cells[0].text = self.eighth_row[0]
        self.eighth_row_cells[0].width = Cm(27.91)
        
        self.ninth_row_cells = self.ninth_row_table.rows[0].cells
        self.ninth_row_cells[0].text = self.ninth_row[0]
        self.ninth_row_cells[0].width = Cm(9.18)
        self.ninth_row_cells[1].text = self.ninth_row[1]
        self.ninth_row_cells[1].width = Cm(9.18)
        self.ninth_row_cells[2].text = self.ninth_row[2]
        self.ninth_row_cells[2].width = Cm(9.19)
        
        self.tenth_row_cells = self.tenth_row_table.rows[0].cells
        self.tenth_row_cells[0].text = self.tenth_row[0]
        self.tenth_row_cells[0].width = Cm(9.18)
        self.tenth_row_cells[1].text = self.tenth_row[1]
        self.tenth_row_cells[1].width = Cm(9.18)
        self.tenth_row_cells[2].text = self.tenth_row[2]
        self.tenth_row_cells[2].width = Cm(9.19)
        
        self.eleventh_row_cells = self.eleventh_row_table.rows[0].cells
        self.eleventh_row_cells[0].text = self.eleventh_row[0]
        self.eleventh_row_cells[0].width = Cm(2.71)
        self.eleventh_row_cells[1].text = self.eleventh_row[1]
        self.eleventh_row_cells[1].width = Cm(6.47)
        self.eleventh_row_cells[2].text = self.eleventh_row[2]
        self.eleventh_row_cells[2].width = Cm(4.59)
        self.eleventh_row_cells[3].text = self.eleventh_row[3]
        self.eleventh_row_cells[3].width = Cm(4.59)
        self.eleventh_row_cells[4].text = self.eleventh_row[4]
        self.eleventh_row_cells[4].width = Cm(4.59)
        self.eleventh_row_cells[5].text = self.eleventh_row[5]
        self.eleventh_row_cells[5].width = Cm(4.60)
        
        self.twelfth_row_cells = self.twelfth_row_table.rows[0].cells
        self.twelfth_row_cells[0].text = self.twelfth_row[0]
        self.twelfth_row_cells[1].text = self.twelfth_row[1]
        self.twelfth_row_cells[2].text = self.twelfth_row[2]
        self.twelfth_row_cells[3].text = self.twelfth_row[3]
        self.twelfth_row_cells[4].text = self.twelfth_row[4]
        self.twelfth_row_cells[5].text = self.twelfth_row[5]
        
        self.thirteenth_row_cells = self.thirteenth_row_table.rows[0].cells
        self.thirteenth_row_cells[0].text = self.thirteenth_row[0]
        self.thirteenth_row_cells[1].text = self.thirteenth_row[1]
        self.thirteenth_row_cells[2].text = self.thirteenth_row[2]
        self.thirteenth_row_cells[3].text = self.thirteenth_row[3]
        self.thirteenth_row_cells[4].text = self.thirteenth_row[4]
        self.thirteenth_row_cells[5].text = self.thirteenth_row[5]
        
        self.fourteenth_row_cells = self.fourteenth_row_table.rows[0].cells
        self.fourteenth_row_cells[0].text = self.fourteenth_row[0]
        self.fourteenth_row_cells[1].text = self.fourteenth_row[1]
        self.fourteenth_row_cells[2].text = self.fourteenth_row[2]
        self.fourteenth_row_cells[3].text = self.fourteenth_row[3]
        self.fourteenth_row_cells[4].text = self.fourteenth_row[4]
        self.fourteenth_row_cells[5].text = self.fourteenth_row[5]
        
        self.fifteenth_row_cells = self.fifteenth_row_table.rows[0].cells
        self.fifteenth_row_cells[0].text = self.fifteenth_row[0]
        self.fifteenth_row_cells[0].width = Cm(2.71)
        self.fifteenth_row_cells[1].text = self.fifteenth_row[1]
        self.fifteenth_row_cells[1].width = Cm(6.47)
        self.fifteenth_row_cells[2].text = self.fifteenth_row[2]
        self.fifteenth_row_cells[2].width = Cm(4.59)
        self.fifteenth_row_cells[3].text = self.fifteenth_row[3]
        self.fifteenth_row_cells[3].width = Cm(4.59)
        self.fifteenth_row_cells[4].text = self.fifteenth_row[4]
        self.fifteenth_row_cells[4].width = Cm(4.59)
        self.fifteenth_row_cells[5].text = self.fifteenth_row[5]
        self.fifteenth_row_cells[5].width = Cm(4.60)
        
    def page_two_fill(self):
        
        # Fills tables for page two depending on input.
        self.page_two_first_row_cells = self.page_two_first_row_table.rows[0].cells
        self.page_two_first_row_cells[0].text = self.page_two_first_row[0]
        self.page_two_first_row_cells[0].width = Cm(20.0)
        self.page_two_first_row_cells[1].text = self.page_two_first_row[1]
        self.page_two_first_row_cells[1].width = Cm(7.91)
        
        self.page_two_second_row_cells = self.page_two_second_row_table.rows[0].cells
        self.page_two_second_row_cells[0].text = self.page_two_second_row[0]
        self.page_two_second_row_cells[0].width = Cm(1.42)
        self.page_two_second_row_cells[1].text = self.page_two_second_row[1]
        self.page_two_second_row_cells[1].width = Cm(1.98)
        self.page_two_second_row_cells[2].text = self.page_two_second_row[2]
        self.page_two_second_row_cells[2].width = Cm(4.88)
        self.page_two_second_row_cells[3].text = self.page_two_second_row[3]
        self.page_two_second_row_cells[3].width = Cm(8.39)
        self.page_two_second_row_cells[4].text = self.page_two_second_row[4]
        self.page_two_second_row_cells[4].width = Cm(6.53)
        self.page_two_second_row_cells[5].text = self.page_two_second_row[5]
        self.page_two_second_row_cells[5].width = Cm(4.72)
        
        self.page_two_third_row_cells = self.page_two_third_row_table.rows[0].cells
        self.page_two_third_row_cells[0].text = self.page_two_third_row[0]
        self.page_two_third_row_cells[0].width = Cm(1.42)
        self.page_two_third_row_cells[1].text = self.page_two_third_row[1]
        self.page_two_third_row_cells[1].width = Cm(1.98)
        self.page_two_third_row_cells[2].text = self.page_two_third_row[2]
        self.page_two_third_row_cells[2].width = Cm(4.88)
        self.page_two_third_row_cells[3].text = self.page_two_third_row[3]
        self.page_two_third_row_cells[3].width = Cm(8.39)
        self.page_two_third_row_cells[4].text = self.page_two_third_row[4]
        self.page_two_third_row_cells[4].width = Cm(6.53)
        self.page_two_third_row_cells[5].text = self.page_two_third_row[5]
        self.page_two_third_row_cells[5].width = Cm(4.72)  
        
        self.page_two_fourth_row_cells = self.page_two_fourth_row_table.rows[0].cells
        self.page_two_fourth_row_cells[0].text = self.page_two_fourth_row[0]
        self.page_two_fourth_row_cells[0].width = Cm(1.42)
        self.page_two_fourth_row_cells[1].text = self.page_two_fourth_row[1]
        self.page_two_fourth_row_cells[1].width = Cm(1.98)
        self.page_two_fourth_row_cells[2].text = self.page_two_fourth_row[2]
        self.page_two_fourth_row_cells[2].width = Cm(4.88)
        self.page_two_fourth_row_cells[3].text = self.page_two_fourth_row[3]
        self.page_two_fourth_row_cells[3].width = Cm(8.39)
        self.page_two_fourth_row_cells[4].text = self.page_two_fourth_row[4]
        self.page_two_fourth_row_cells[4].width = Cm(6.53)
        self.page_two_fourth_row_cells[5].text = self.page_two_fourth_row[5]
        self.page_two_fourth_row_cells[5].width = Cm(4.72)
        
        self.page_two_fifth_row_cells = self.page_two_fifth_row_table.rows[0].cells
        self.page_two_fifth_row_cells[0].text = self.page_two_fifth_row[0]
        self.page_two_fifth_row_cells[0].width = Cm(1.42)
        self.page_two_fifth_row_cells[1].text = self.page_two_fifth_row[1]
        self.page_two_fifth_row_cells[1].width = Cm(1.98)
        self.page_two_fifth_row_cells[2].text = self.page_two_fifth_row[2]
        self.page_two_fifth_row_cells[2].width = Cm(4.88)
        self.page_two_fifth_row_cells[3].text = self.page_two_fifth_row[3]
        self.page_two_fifth_row_cells[3].width = Cm(8.39)
        self.page_two_fifth_row_cells[4].text = self.page_two_fifth_row[4]
        self.page_two_fifth_row_cells[4].width = Cm(6.53)
        self.page_two_fifth_row_cells[5].text = self.page_two_fifth_row[5]
        self.page_two_fifth_row_cells[5].width = Cm(4.72)        
          
    def savefile(self):
        
        # Saves the file with specified name and location.
        self.document.save(self.filename)
